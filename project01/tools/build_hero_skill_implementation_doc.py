from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "design" / "first-batch-8-heroes-latest-skills.md"
CONFIG = ROOT / "docs" / "design" / "first-batch-8-heroes-y3-skill-table-v0.md"
OUT = ROOT / "deliverables" / "三国斗阵_Y3_英雄技能实现说明_v0.2.docx"

FONT = "Microsoft YaHei"
INK = "222222"
MUTED = "666666"
BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F6F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(sum(widths_cm) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_cm):
        grid_col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=10, bold=False, color=INK) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def write_cell(cell, text: str, size=8.5, bold=False, color=INK, center=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margin(cell)


def parse_config() -> dict[tuple[str, str], dict[str, str]]:
    rows = {}
    for line in CONFIG.read_text(encoding="utf-8").splitlines():
        if not line.startswith("|") or "---" in line or "hero |" in line:
            continue
        cells = [part.strip() for part in line.strip().strip("|").split("|")]
        if len(cells) != 11:
            continue
        hero, skill_id, skill, kind, unlock, unit_type, target, card, pos, ai, manual = cells
        rows[(hero, skill)] = {
            "skill_id": skill_id,
            "target": target,
            "card": card,
            "unit_type": unit_type,
        }
    return rows


def parse_existing_heroes() -> dict[str, dict]:
    text = SOURCE.read_text(encoding="utf-8")
    sections = re.split(r"^## ", text, flags=re.MULTILINE)
    heroes = {}
    for section in sections:
        if not section:
            continue
        name = section.splitlines()[0].strip()
        if name not in {"黄忠", "刘备", "吕蒙", "赵云", "张飞", "夏侯惇", "吕布", "诸葛亮"}:
            continue
        positioning = re.search(r"\*\*定位：\*\*\s*(.+)", section)
        duty = re.search(r"\*\*核心职责：\*\*\s*(.+)", section)
        loop = re.search(r"\*\*专属循环：([^*]+)\*\*", section)
        loop_name = loop.group(1).strip() if loop else "-"
        before_table = section.split("| 阶段 |", 1)[0]
        loop_lines = re.findall(r"^-\s+(.+)$", before_table, flags=re.MULTILINE)
        table_lines = [line for line in section.splitlines() if line.startswith("| 阶段") or (line.startswith("|") and "---" not in line)]
        skills = []
        for line in table_lines[1:]:
            cells = [part.strip() for part in line.strip().strip("|").split("|")]
            if len(cells) == 4:
                skills.append({"stage": cells[0], "name": cells[1], "type": cells[2], "effect": cells[3]})
        heroes[name] = {
            "positioning": positioning.group(1).strip() if positioning else "-",
            "duty": duty.group(1).strip() if duty else "-",
            "loop_name": loop_name,
            "loop_lines": loop_lines,
            "skills": skills,
        }
    return heroes


EXTRA_NOTES = {
    "黄忠": [
        "后排条件只读取本波次战前锁定站位 `7/8/9`；战斗移动不会清空蓄势。",
        "蓄势最多 3 层；受到近战伤害时清空。",
        "`穿云贯阵` 只作用当前一路，不跨路；`百步穿杨` 可打断队长或精英的技能释放。",
        "`箭阵老将` 不改变远程小兵目标，只调整其攻速和对小兵伤害。",
    ],
    "刘备": [
        "刘备的主动技能只筛选友方队长，永不治疗或施加护盾给小兵。",
        "仁德：友方队长生命低于 40% 时触发；每名队长每 8 秒最多提供 1 层；最多 3 层；波次结束清空。",
        "`仁望归心` 不消耗仁德，仅在刘备释放主动后检查是否仍持有仁德以延长恢复。",
        "`汉室余泽` 的 12 秒内置冷却属于刘备自身，不按每个受益队长独立计时。",
    ],
    "吕蒙": [
        "背刺：每次攻击命中时，实时检测吕蒙是否位于目标背后 120 度范围内。",
        "隐身期间可普攻且普攻不破隐；释放主动、受到伤害或持续时间结束时解除隐身。",
        "隐身只降低普通敌人的选敌优先级；Boss、必中技能和范围技能不受影响。",
        "`双刀断影` 固定造成 2 次斩击，每次独立进行背刺判定。侧翼条件读取战前锁定侧翼。",
    ],
    "赵云": [
        "游龙：成功闪避时获得 1 层，最多 3 层；不存在旧版“闪身次数”机制。",
        "`游龙身法` 固定提供闪避 +15%；`白马轻骑` 的侧翼条件读取战前锁定侧翼。",
        "`一身是胆`：消耗游龙后获得闪避 +7%、普攻速度 +15%，持续 4 秒；持有 3 层游龙时普攻速度 +20%。",
    ],
    "张飞": [
        "守势：每次受到近战伤害获得 1 层，每 1 秒最多 1 层，最多 3 层；每层近战减伤 6%；持续 5 秒，重复获得刷新持续时间。",
        "技能消耗守势时，同时失去该层对应的近战减伤。前排条件读取战前锁定前排 `1/2/3`。",
        "`万夫莫开`：生命低于 40% 时立即获得 3 层守势，12 秒内置冷却。",
    ],
    "夏侯惇": [
        "刚烈：受到近战伤害时获得 1 层，每 1 秒最多 1 层，最多 4 层；护盾吸收的近战伤害同样计入。",
        "4 层刚烈自动消耗并反斩当前目标区域，反斩内置冷却 3 秒；远程、法术与持续伤害不提供刚烈。",
        "前排条件读取战前锁定前排 `1/2/3`。",
    ],
    "吕布": [
        "威势：参与击破敌方队长或精英时获得 1 层；“参与”指死亡前 3 秒内吕布造成过伤害。",
        "威势最多 3 层，每层攻击 +6%，当前波次内保留；小兵与召唤物不提供威势。",
        "`阵前叫战` 是进攻挑战口径；`飞将压迫` 只降低当前目标区域内敌方小兵攻击。前排条件读取战前锁定前排。",
    ],
    "诸葛亮": [
        "谋略：战前锁定在后排 `7/8/9` 时，每 4 秒获得 1 层，最多 3 层；受到近战伤害时失去 1 层。战斗移动不会清空谋略。",
        "`八卦阵`、`八阵定军` 为当前目标区域技能；`借东风` 为当前一路技能，均不跨路。",
        "眩晕、打断等控制对 Boss 按 Boss 抗性和关键阶段免疫规则结算。",
    ],
}


NEW_HEROES = {
    "司马懿": {
        "positioning": "敌方队长削弱、布置与收束爆发的法术远程英雄。",
        "duty": "远程魔法书；远程小兵；职业：法。",
        "loop_name": "筹谋 / 定局",
        "loop_lines": [
            "每释放 1 次主动技能，获得 1 层筹谋，最多 3 层。",
            "达到 3 层后进入“定局”状态：下一次主动技能获得强化并消耗全部筹谋。",
            "筹谋仅在当前波次内有效，波次结束清空。",
        ],
        "skills": [
            {"stage": "阶段1", "name": "鹰视", "type": "主动", "effect": "指定一名敌方队长，造成 240% 法术伤害，并使其在 6 秒内受到司马懿造成的伤害提高 12%。处于“定局”时，伤害提高至 330%，增伤提高至 18%。"},
            {"stage": "阶段1", "name": "深谋远虑", "type": "被动", "effect": "每释放 1 次主动技能，获得 1 层筹谋。筹谋最多 3 层。"},
            {"stage": "阶段2", "name": "缚心", "type": "主动", "effect": "指定一名敌方队长，造成 120% 法术伤害，使其在 5 秒内造成的伤害降低 20%，攻击速度降低 25%。处于“定局”时，持续时间延长至 7 秒，伤害降低提高至 30%，攻击速度降低提高至 35%。"},
            {"stage": "阶段2", "name": "隐忍待变", "type": "被动", "effect": "司马懿首次低于 45% 最大生命时，获得可吸收自身 25% 最大生命伤害的护盾，持续 8 秒。每波次最多触发 1 次。"},
            {"stage": "阶段3", "name": "冢虎决", "type": "主动", "effect": "指定一名敌方队长，造成 300% 法术伤害。若目标正处于“鹰视”或“缚心”效果下，伤害提高至 450%。处于“定局”时，基础伤害提高至 420%，满足条件时提高至 600%。"},
            {"stage": "阶段3", "name": "将略统军", "type": "被动", "effect": "司马懿进入“定局”状态时，其小队远程小兵在 6 秒内对正在攻击的敌方队长造成的伤害提高 25%。"},
        ],
        "notes": [
            "三个主动技能均可进入卡牌池；关闭司马懿自动释放时，已解锁主动技能全部进入公共卡池。",
            "`将略统军` 不改变远程小兵目标，仅在其自然攻击敌方队长时提高伤害。",
            "主动技能冷却尚未配置；应在数值表中补齐。",
        ],
    },
    "甘宁": {
        "positioning": "以普攻装填资源，压制与斩杀关键队长的远程贼系英雄。",
        "duty": "远程双火枪；近战小兵；职业：贼。",
        "loop_name": "火药 / 满膛",
        "loop_lines": [
            "每以普通攻击命中敌方单位 3 次，获得 1 层火药，最多 3 层。",
            "达到 3 层后进入“满膛”状态：下一次主动技能获得强化并消耗全部火药。",
            "火药仅在当前波次内有效，波次结束清空。",
        ],
        "skills": [
            {"stage": "阶段1", "name": "锦帆连铳", "type": "主动", "effect": "指定一名敌方队长，连续射击 3 次，每次造成 85% 物理伤害。处于“满膛”时，改为连续射击 5 次，每次造成 95% 物理伤害。"},
            {"stage": "阶段1", "name": "火药装填", "type": "被动", "effect": "每以普通攻击命中敌方单位 3 次，获得 1 层火药。火药最多 3 层。"},
            {"stage": "阶段2", "name": "震胆铳击", "type": "主动", "effect": "指定一名敌方队长，造成 180% 物理伤害，使其在 5 秒内造成的伤害降低 20%，攻击速度降低 25%。处于“满膛”时，伤害提高至 250%，持续时间延长至 7 秒，伤害降低提高至 30%，攻击速度降低提高至 35%。"},
            {"stage": "阶段2", "name": "锦帆侧击", "type": "被动", "effect": "若甘宁在战前布阵时位于侧翼，闪避提高 8%，攻击速度提高 15%。该效果在本波次内固定生效。"},
            {"stage": "阶段3", "name": "斩旗夺将", "type": "主动", "effect": "指定一名敌方队长，造成 250% 物理伤害；若其当前生命低于 35%，改为造成 450% 物理伤害。处于“满膛”时，生命判断提高至 50%，伤害提高至 540%。"},
            {"stage": "阶段3", "name": "短兵接舷", "type": "被动", "effect": "甘宁的主动技能命中敌方队长后，其小队近战小兵在 6 秒内对该队长造成的伤害提高 25%，但不会改变小兵原本的攻击目标。"},
        ],
        "notes": [
            "三个主动技能均可进入卡牌池；关闭甘宁自动释放时，已解锁主动技能全部进入公共卡池。",
            "侧翼条件只读取战前锁定的 `1/4/7` 或 `3/6/9`，战斗移动不改变效果。",
            "`短兵接舷` 只给甘宁小队近战小兵的自然攻击加成，禁止强制切换目标。",
            "主动技能冷却尚未配置；应在数值表中补齐。",
        ],
    },
}


ROLE = {
    "吕布": "战", "张飞": "战", "吕蒙": "贼", "黄忠": "射", "诸葛亮": "法",
    "赵云": "战", "刘备": "牧", "夏侯惇": "战", "司马懿": "法", "甘宁": "贼",
}

CONFIG_PREFIX = {
    "黄忠": "huangzhong", "刘备": "liubei", "吕蒙": "lvmeng", "赵云": "zhaoyun",
    "张飞": "zhangfei", "夏侯惇": "xiahoudun", "吕布": "lvbu", "诸葛亮": "zhugeliang",
    "司马懿": "simayi", "甘宁": "ganning",
}


def target_label(raw: str) -> str:
    if not raw:
        return "待定义"
    return raw.replace("_", " ")


def add_heading(doc, text: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_font(run, size=15 if level == 1 else 10.5, bold=True, color=BLUE)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=9)


def add_hero(doc, hero_name: str, hero: dict, config: dict) -> None:
    add_heading(doc, f"{hero_name}  |  职业：{ROLE[hero_name]}")
    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    set_table_widths(meta, [4.7, 18.0])
    meta_rows = [
        ("定位", hero["positioning"]),
        ("实现前缀", CONFIG_PREFIX[hero_name]),
    ]
    for row, (label, value) in zip(meta.rows, meta_rows):
        write_cell(row.cells[0], label, size=8.5, bold=True, color=BLUE)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        write_cell(row.cells[1], value, size=8.5)

    add_heading(doc, f"核心机制：{hero['loop_name']}", level=2)
    callout = doc.add_table(rows=1, cols=1)
    set_table_widths(callout, [22.7])
    set_cell_shading(callout.cell(0, 0), LIGHT_GRAY)
    write_cell(callout.cell(0, 0), "\n".join(hero["loop_lines"]), size=9)

    add_heading(doc, "技能实现", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["阶段", "类型", "技能", "技能 ID", "目标 / 卡池", "效果与实现规则"]
    widths = [1.25, 1.1, 2.2, 4.2, 3.1, 10.85]
    set_table_widths(table, widths)
    for cell, label in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        write_cell(cell, label, size=8.3, bold=True, color=BLUE, center=True)

    for skill in hero["skills"]:
        row = table.add_row()
        config_row = config.get((hero_name, skill["name"]), {})
        skill_id = config_row.get("skill_id", f"{CONFIG_PREFIX[hero_name]}_{skill['name']}")
        card = "进入卡池" if skill["type"] == "主动" else "不进卡池"
        target = target_label(config_row.get("target", "指定队长" if skill["type"] == "主动" else "被动触发"))
        values = [skill["stage"], skill["type"], skill["name"], skill_id, f"{target}\n{card}", skill["effect"]]
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            write_cell(cell, value, size=7.8 if index in (3, 4, 5) else 8.2, center=index in (0, 1))

    notes = hero.get("notes") or EXTRA_NOTES.get(hero_name, [])
    if notes:
        add_heading(doc, "特殊实现约束", level=2)
        for note in notes:
            add_bullet(doc, note)


def build() -> None:
    existing = parse_existing_heroes()
    config = parse_config()
    heroes = {**existing, **NEW_HEROES}
    order = ["吕布", "张飞", "吕蒙", "黄忠", "诸葛亮", "赵云", "刘备", "夏侯惇", "司马懿", "甘宁"]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("三国斗阵 | Y3 英雄技能实现说明")
    set_font(header_run, size=8, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("内部工作文档 | v0.2 | 2026-07-10")
    set_font(footer_run, size=8, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(3)
    set_font(title.add_run("《三国斗阵》Y3 英雄技能实现说明"), size=22, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(9)
    set_font(subtitle.add_run("首批十名英雄 | 程序实现参考 | v0.2"), size=10.5, color=MUTED)
    note = doc.add_table(rows=1, cols=1)
    set_table_widths(note, [22.7])
    set_cell_shading(note.cell(0, 0), LIGHT_GRAY)
    write_cell(note.cell(0, 0), "本文件仅包含英雄技能实现内容。已确认的触发、目标、次数、阈值与内置冷却按当前设计写入；源设计尚未定义的技能伤害、冷却、施法距离等，统一标记为“待数值配置”，不得由程序侧自行推定。", size=9)

    for index, hero_name in enumerate(order):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_hero(doc, hero_name, heroes[hero_name], config)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "三国斗阵 Y3 英雄技能实现说明 v0.2"
    doc.core_properties.subject = "首批十名英雄技能实现参考"
    doc.core_properties.author = "project01"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
