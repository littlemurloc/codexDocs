from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "三国斗阵-军令效果设计清单-V1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "5A677A"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
FONT = "Microsoft YaHei"
CONTENT_WIDTH = 9360


def set_run_font(run, size=11, color=INK, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths):
        grid_cols[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_paragraph_border_bottom(paragraph, color="D5DCE6", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=8.5, color=MUTED)


def add_text(doc, text, style=None, before=0, after=6, color=INK, size=11, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    for run in p.runs:
        set_run_font(run, size=10.5)
    if not p.runs:
        set_run_font(p.add_run(text), size=10.5)
    else:
        p.runs[0].text = text
    return p


def add_label_value(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], HEADER_FILL)
        p1 = cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=9.5, color=DARK_BLUE, bold=True)
        p2 = cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_card(doc, name, card_type, text, meta=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(name)
    set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    tag = p.add_run(f"  {card_type}")
    set_run_font(tag, size=9.5, color=BLUE, bold=True)
    ident = p.add_run(f"  ID: {CARD_IDS[name]}")
    set_run_font(ident, size=8.5, color=MUTED)
    if meta:
        m = p.add_run(f"  {meta}")
        set_run_font(m, size=9.5, color=MUTED)
    body = doc.add_paragraph()
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(5)
    body.paragraph_format.line_spacing = 1.25
    run = body.add_run(text)
    set_run_font(run, size=10.5)


def add_result(doc, name, card_type, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"自动结果：{name}")
    set_run_font(r, size=12, color=BLUE, bold=True)
    tag = p.add_run(f"  {card_type}")
    set_run_font(tag, size=9.5, color=DARK_BLUE, bold=True)
    ident = p.add_run(f"  ID: {CARD_IDS[name]}")
    set_run_font(ident, size=8.5, color=MUTED)
    for label, value in lines:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.line_spacing = 1.22
        a = p2.add_run(f"{label}：")
        set_run_font(a, size=10.5, color=DARK_BLUE, bold=True)
        b = p2.add_run(value)
        set_run_font(b, size=10.5)


SETS = [
    {
        "name": "破军 I",
        "mode": "2 被动 + 1 主动，主动进阶",
        "unlock": "虎贲 II、飞骑 II",
        "cards": [
            ("锋矢军略", "被动", "全部友方队长攻击提高 8%；其小兵攻击提高 4%。"),
            ("战鼓军略", "被动", "每当敌方队长或精英死亡时，全部存活友方队长及小兵攻击速度提高 4%，最多叠加 3 层。本波内生效。普通小兵死亡不触发。"),
            ("斩将令", "主动", "选择 1 名敌方队长或精英。0.6 秒预警后，连续造成 3 次伤害，每次为目标最大生命值 2% 的真实伤害，共 6%。目标在预警期死亡或离场则返还本次使用；目标在三次伤害中死亡，剩余伤害取消且不改选目标。"),
        ],
        "result": ("破阵令", "主动进阶", [
            ("继承", "斩将令自动替换为破阵令，保留原军令库与 HUD 槽位；锋矢军略与战鼓军略继续生效。"),
            ("效果", "选择敌方战前上、中、下任一路。0.8 秒预警后，该路全部存活敌方单位各受到自身最大生命值 4% 的真实伤害；该路队长与精英额外受到自身最大生命值 2% 的真实伤害。全部命中单位防御降低 15%，持续 4 秒。"),
            ("失效", "目标路在结算前已无存活敌人时返还本次使用。击退表现不改变战前路、行列或站位属性。"),
            ("解锁", "虎贲 II、飞骑 II。"),
        ]),
    },
    {
        "name": "虎贲 II",
        "mode": "3 被动，后台融合",
        "unlock": "无第三层解锁",
        "cards": [
            ("铁甲军略", "被动", "战前位于前排 123 的友方队长及其小兵防御提高 12%。"),
            ("盾列军略", "被动", "战前位于前排 123 的友方队长及其小兵最大生命值提高 10%。"),
            ("死守军略", "被动", "每个前排友方队伍每波首次由其队长生命值降至 40% 或以下时，该队长及其存活小兵受到伤害降低 20%，持续 6 秒。"),
        ],
        "result": ("固阵军略", "后台被动融合", [
            ("继承", "保留三张组件效果。"),
            ("新增", "只要至少 1 名战前位于前排的友方队长仍存活，所有战前位于中排或后排的友方单位受到伤害降低 5%。最后一名前排队长死亡后，该效果结束。"),
        ]),
    },
    {
        "name": "飞骑 II",
        "mode": "2 被动 + 1 主动，主动进阶",
        "unlock": "无第三层解锁",
        "cards": [
            ("良马军略", "被动", "战前位于侧翼 147 或 369 的友方队长及其小兵移动速度提高 12%，闪避提高 8%。"),
            ("侧击军略", "被动", "战前位于侧翼的友方队伍，攻击战前位于中排或后排的敌方队长及其小兵时，造成伤害提高 10%。不改变选敌。"),
            ("奇袭令", "主动", "选择敌方战前上、中、下任一路。0.6 秒预警后，锁定该路战前列最靠后的存活敌方队伍，优先级为后排、中排、前排。目标队伍中的队长及所有存活小兵，各受到自身最大生命值 5% 的真实伤害。目标路在结算前清空则返还本次使用。"),
        ],
        "result": ("奔袭令", "主动进阶", [
            ("继承", "奇袭令自动替换，保留两张被动效果。"),
            ("效果", "选择敌方战前上、中、下任一路。0.8 秒预警后，按后排、中排、前排优先级锁定该路最靠后的存活敌方队伍。该队伍中全部存活单位各受到自身最大生命值 8% 的真实伤害，攻击速度降低 35%、移动速度降低 30%，持续 5 秒。"),
            ("表现", "军令召来的飞骑突击，不移动已有英雄队伍。"),
        ]),
    },
    {
        "name": "火计 I",
        "mode": "2 被动 + 1 主动，主动进阶",
        "unlock": "连环 II、伏兵 II",
        "cards": [
            ("火油军略", "被动", "每波开始时，敌方战前位于前排 123 的队长及其小兵获得灼烧，持续 4 秒，每秒造成自身最大生命值 0.6% 的真实伤害，共 4 次。"),
            ("风助军略", "被动", "带有灼烧的任意敌方单位死亡时，对其所在敌方路的其余存活单位各造成自身最大生命值 2% 的真实伤害。每条敌方路每波最多触发 1 次。该伤害不附加灼烧。"),
            ("焚营令", "主动", "选择敌方战前上、中、下任一路。0.6 秒预警后，该路全部存活敌方单位获得灼烧，持续 5 秒，每秒造成自身最大生命值 1% 的真实伤害，共 5 次。目标路在结算前清空则返还本次使用。"),
        ],
        "result": ("火攻令", "主动进阶", [
            ("继承", "焚营令自动替换，保留两张被动效果。"),
            ("效果", "选择敌方战前上、中、下任一路。0.8 秒预警后，该路全部存活敌方单位获得灼烧，持续 5 秒，每秒造成自身最大生命值 1.5% 的真实伤害，共 5 次。灼烧期间，该路单位造成的伤害降低 15%。"),
            ("解锁", "连环 II、伏兵 II。"),
        ]),
    },
    {
        "name": "连环 II",
        "mode": "2 被动 + 1 主动，主动进阶",
        "unlock": "无第三层解锁",
        "cards": [
            ("缚舟军略", "被动", "处于灼烧状态且生命值低于 40% 的敌方队长或精英，受到的所有伤害提高 12%。"),
            ("断索军略", "被动", "带有灼烧的敌方队长或精英死亡时，其所在敌方路的其余存活单位防御降低 12%，持续 5 秒。每条敌方路每波最多触发 1 次。"),
            ("连环令", "主动", "选择敌方战前上、中、下任一路；该路没有存活队长或精英时不可选择。0.6 秒预警后，锁定该路生命比例最低的敌方队长或精英，并在选择阶段高亮预览。对其造成最大生命值 4% 的真实伤害，施加持续 5 秒、每秒最大生命值 1% 真实伤害的灼烧，并施加持续 5 秒的火媒。火媒期间目标死亡时，相邻敌路全部存活单位各受到最大生命值 2.5% 的真实伤害。连锁伤害不附加灼烧。"),
        ],
        "result": ("焚舟令", "主动进阶", [
            ("继承", "连环令自动替换，保留两张被动效果。"),
            ("效果", "选择敌方战前上、中、下任一路；按相同规则锁定该路生命比例最低的存活敌方队长或精英。0.8 秒预警后，对其造成最大生命值 6% 的真实伤害，施加持续 5 秒、每秒最大生命值 1.5% 真实伤害的灼烧，并施加持续 8 秒的火媒。"),
            ("火媒死亡", "相邻敌路全部存活单位各受到最大生命值 4% 的真实伤害，且防御降低 12%，持续 4 秒。连锁伤害不附加灼烧。"),
            ("失效", "目标在预警阶段死亡或离场则返还本次使用。"),
        ]),
    },
    {
        "name": "伏兵 II",
        "mode": "2 被动 + 1 主动，主动进阶",
        "unlock": "无第三层解锁",
        "cards": [
            ("耳目军略", "被动", "每波开始后的 6 秒内，全部敌方队长与精英攻击速度降低 10%。小兵不受影响。"),
            ("断粮军略", "被动", "每名敌方队长或精英每波首次生命值降至 50% 或以下时，其造成的伤害降低 12%，持续至本波结束。"),
            ("伏兵令", "主动", "选择敌方战前上、中、下任一路。0.6 秒预警后在该路布置伏兵，2 秒后发动突袭。该路全部存活敌方单位各受到最大生命值 2% 的真实伤害，攻击速度和移动速度降低 30%，持续 4 秒；该路队长与精英造成的伤害额外降低 15%，持续 4 秒。突袭发动前目标路清空则返还本次使用。"),
        ],
        "result": ("十面埋伏令", "主动进阶", [
            ("继承", "伏兵令自动替换，保留两张被动效果。"),
            ("效果", "选择敌方战前上、中、下任一路。0.6 秒预警后布置伏兵，1.5 秒后发动突袭。该路全部存活敌方单位各受到最大生命值 4% 的真实伤害，攻击速度和移动速度降低 40%，持续 5 秒；该路队长与精英造成的伤害降低 20%，持续 5 秒。"),
            ("表现", "军令召来的伏兵或箭阵，不移动已有英雄队伍。"),
        ]),
    },
    {
        "name": "屯田 I",
        "mode": "3 被动，后台融合",
        "unlock": "募兵 II、军械 II",
        "cards": [
            ("屯垦军略", "被动", "每次进入备战阶段时获得 1 点军资；获得时机在报价出现前。"),
            ("节用军略", "被动", "每个备战阶段的第一次刷新报价免费；后续刷新遵循正常价格规则。"),
            ("缴获军略", "被动", "每波中首次击败敌方精英或 Boss 时，获得 1 点军资。普通敌方队长不触发。"),
        ],
        "result": ("富国军略", "后台被动融合", [
            ("继承", "保留三张组件效果。"),
            ("新增", "每波进入战斗前，按备战结束时持有的军资获得本波储备收益，军资不被消耗。持有 5 点或以上军资时，全部友方队长及小兵最大生命值提高 8%；持有 8 点或以上军资时，改为最大生命值提高 12%，并额外攻击提高 8%。"),
            ("解锁", "募兵 II、军械 II。"),
        ]),
    },
    {
        "name": "募兵 II",
        "mode": "3 被动，后台融合",
        "unlock": "无第三层解锁",
        "cards": [
            ("募卒军略", "被动", "每个备战阶段，可消耗 1 点军资，为 1 名仍存活的友方队长补充 1 名小兵。不得超过该队长当前升星对应的小兵上限；每个备战阶段最多补充 1 名。"),
            ("练卒军略", "被动", "所有友方小兵最大生命值提高 12%，攻击提高 8%。对战斗开始时已有小兵及后续补充的小兵都生效。"),
            ("成阵军略", "被动", "每波开始时，若一名友方队长的小兵数量达到其个人上限，该队长及其小兵攻击速度提高 10%、受到伤害降低 6%，持续本波。战斗中小兵阵亡不会移除本波已获得的效果。"),
        ],
        "result": ("精兵强军", "后台被动融合", [
            ("继承", "保留三张组件效果。"),
            ("新增", "每个备战阶段最多可补充 2 名小兵；第 1 次补员免费，第 2 次补员消耗 1 点军资。满编队伍的本波效果改为攻击速度提高 15%、受到伤害降低 10%，并额外攻击提高 8%。"),
            ("边界", "队长死亡后，小兵仍按既定规则溃败离场；本军略不改变队长与小兵的绑定关系。"),
        ]),
    },
    {
        "name": "军械 II",
        "mode": "3 被动，后台融合",
        "unlock": "无第三层解锁",
        "cards": [
            ("精锻军略", "被动", "所有友方队长及小兵攻击提高 8%。"),
            ("甲胄军略", "被动", "所有友方队长及小兵防御提高 10%。"),
            ("校准军略", "被动", "所有友方队长及小兵攻击速度提高 10%，普攻暴击率提高 5%。"),
        ],
        "result": ("百工军略", "后台被动融合", [
            ("继承", "保留三张组件效果。"),
            ("新增", "每名友方单位独立计算普通攻击次数。每波中第 4、8、12 等第 4 次倍数的普通攻击无视目标 30% 防御。攻击次数在每波开始时重置；技能伤害、军令伤害和持续伤害不参与计数。"),
        ]),
    },
]

INDEPENDENTS = [
    ("鼓舞令", "主动", "选择我方战前上、中、下任一路。0.3 秒后，该路所属队长及小兵攻击提高 15%、攻击速度提高 20%，持续 6 秒。", "开局候选"),
    ("固守令", "主动", "选择我方战前上、中、下任一路。0.3 秒后，该路所属队长及小兵获得相当于各自最大生命值 10% 的护盾，护盾最多持续 6 秒。", "开局候选"),
    ("体魄军略", "被动", "所有友方队长及小兵最大生命值提高 6%。", "开局候选"),
    ("锐气军略", "被动", "每波战斗开始后的 8 秒内，所有友方队长及小兵攻击提高 10%。", "开局候选"),
    ("行军军略", "被动", "所有友方队长及小兵移动速度提高 10%、闪避提高 4%。", "第 3 波起候选"),
    ("保全军略", "被动", "每波胜利结算时，若本波没有友方队长死亡，获得 1 点军资；每波最多触发 1 次，小兵阵亡不影响判定。", "第 3 波起候选"),
]


SET_IDS = {
    "破军 I": "break_i",
    "虎贲 II": "tiger_ii",
    "飞骑 II": "flying_ii",
    "火计 I": "fire_i",
    "连环 II": "chain_ii",
    "伏兵 II": "ambush_ii",
    "屯田 I": "farm_i",
    "募兵 II": "recruit_ii",
    "军械 II": "armory_ii",
}

CARD_IDS = {
    "锋矢军略": "cmd_break_spearhead", "战鼓军略": "cmd_break_war_drum", "斩将令": "cmd_break_execute", "破阵令": "cmd_break_formation",
    "铁甲军略": "cmd_tiger_iron_armor", "盾列军略": "cmd_tiger_shield_line", "死守军略": "cmd_tiger_hold_fast", "固阵军略": "cmd_tiger_hold_formation",
    "良马军略": "cmd_flying_good_mount", "侧击军略": "cmd_flying_flank_assault", "奇袭令": "cmd_flying_raid", "奔袭令": "cmd_flying_charge",
    "火油军略": "cmd_fire_oil", "风助军略": "cmd_fire_wind", "焚营令": "cmd_fire_burn_camp", "火攻令": "cmd_fire_attack",
    "缚舟军略": "cmd_chain_bind_boats", "断索军略": "cmd_chain_cut_rope", "连环令": "cmd_chain_link_order", "焚舟令": "cmd_chain_burn_boats",
    "耳目军略": "cmd_ambush_scout", "断粮军略": "cmd_ambush_cut_supply", "伏兵令": "cmd_ambush_order", "十面埋伏令": "cmd_ambush_ten_sides",
    "屯垦军略": "cmd_farm_reclaim", "节用军略": "cmd_farm_thrift", "缴获军略": "cmd_farm_spoils", "富国军略": "cmd_farm_prosperity",
    "募卒军略": "cmd_recruit_enlist", "练卒军略": "cmd_recruit_train", "成阵军略": "cmd_recruit_formation", "精兵强军": "cmd_recruit_elite_force",
    "精锻军略": "cmd_armory_forging", "甲胄军略": "cmd_armory_armor", "校准军略": "cmd_armory_calibration", "百工军略": "cmd_armory_masterwork",
    "鼓舞令": "cmd_independent_inspire", "固守令": "cmd_independent_hold", "体魄军略": "cmd_independent_vitality", "行军军略": "cmd_independent_march", "锐气军略": "cmd_independent_vigor", "保全军略": "cmd_independent_preserve",
}


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("三国斗阵 | 军令效果设计清单 V1")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(p)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("军令效果设计清单 V1")
    set_run_font(r, size=24, color=INK, bold=True)
    set_paragraph_border_bottom(p, color="C7D4E5", size="10")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("《三国斗阵》 | 程序实现参考 | 2026-07-17")
    set_run_font(r, size=11, color=MUTED)

    add_label_value(doc, [
        ("文档用途", "首批军令卡、Set 融合、市场与运行时效果的程序实现参考。"),
        ("内容范围", "27 张 Set 组件、9 个自动结果、6 张独立军略，以及全局目标、军令点、补员与灼烧规则。"),
        ("当前状态", "效果与流程已确认；基础军资、报价权重、属性叠加公式仍待数值核算。"),
        ("权威来源", "docs/design/tactical-command-card-system-v1.md 与 tactical-command-card-catalog-v1.md。"),
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    r = p.add_run("实现边界：本文件不定义英雄技能、卡牌价格、初始军资、刷新价格曲线、市场权重或完整属性公式。相关系统应保持可配置，避免把当前效果数值写死在流程代码中。")
    set_run_font(r, size=10, color=DARK_BLUE, bold=True)


def add_global_rules(doc):
    doc.add_heading("1. 全局运行契约", level=1)
    add_text(doc, "以下规则先于逐卡效果执行。", after=6)
    add_label_value(doc, [
        ("单局范围", "军令卡、军资、Set 进度、融合结果和相关属性收益均只在当前单局存在，结算后清空。"),
        ("主动军令", "备战时从已拥有主动军令中选择至多 3 张进入 HUD。普通/精英波 2 点军令点，Boss 波 3 点；每张主动军令每波最多使用 1 次，默认消耗 1 点。未用点数波末清空。"),
        ("战前路归属", "军令中的上、中、下路均为开战前归属，不读取实时地图坐标。敌方上、中、下路为 147 / 258 / 369；队长与其小兵共享归属。"),
        ("目标选择", "一张主动军令最多一次手动选择。系统自动锁定时须预览最终对象；不得要求玩家先选友方再选敌方。"),
        ("小兵损失", "小兵阵亡跨波次保留损失，默认不自动补全。只有补员军略能在备战补兵，且不能超过队长当前升星的小兵上限。"),
    ])

    doc.add_heading("2. 通用状态与结算", level=1)
    add_label_value(doc, [
        ("灼烧", "每秒结算一次，按目标最大生命值百分比造成真实伤害。同一单位不叠加多个灼烧：较高每秒伤害替换现有效果；相同或较低伤害只把持续时间延长至较长值。"),
        ("护盾", "以目标自身最大生命值百分比计算；持续时间和是否提前破裂以具体军令为准。"),
        ("预警与返还", "军令在预警后结算。卡牌声明返还时，目标或目标路在结算前完全失效则不消耗军令点，并恢复本波可用状态。"),
        ("属性修正", "攻击、防御、生命、攻速、移速、闪避、暴击、造成伤害和受到伤害均应走统一可配置 Modifier 管线。叠加、舍入和结算顺序待战斗数值层确定。"),
        ("临时表现", "飞骑、伏兵等表现是军令召来的临时战术事件，不移动既有英雄队伍，不改变其战前路、行列和站位属性。"),
    ])

    doc.add_heading("3. 市场、融合与解锁", level=1)
    add_bullet(doc, "每次备战展示 3 个报价。玩家可购买多张报价卡，也可消耗军资刷新整组报价。")
    add_bullet(doc, "报价由路线位、探索位和自由位组成。路线位优先服务已开始但未完成的 Set；探索位服务未开始根系、独立军略和已解锁二阶路线；自由位从全部合法候选中产生。")
    add_bullet(doc, "开局候选为三条根系的 9 张组件，加上鼓舞令、固守令、体魄军略、锐气军略。第 3 波起加入行军军略和保全军略。")
    add_bullet(doc, "根系完成后，相关两条二阶路线进入候选池。后续连续两次备战报价中，各保证出现 1 张新解锁组件，分别来自两条二阶路线，顺序随机，仍需正常购买。")
    add_bullet(doc, "2 被动 + 1 主动的 Set 使基础主动自动进阶并保留原 HUD 槽位；3 被动的 Set 自动吞噬为后台高阶军略。高阶结果不进入市场，不额外收费。")


def add_set_overview(doc):
    doc.add_heading("4. Set 图谱总览", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 2200, 2800, 2460])
    headers = ["Set", "融合方式", "自动结果", "解锁"]
    for i, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])
    for entry in SETS:
        result_name, _, _ = entry["result"]
        row = table.add_row().cells
        values = [entry["name"], entry["mode"], result_name, entry["unlock"]]
        for i, value in enumerate(values):
            p = row[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_sets(doc):
    doc.add_heading("5. Set 组件与自动结果", level=1)
    for entry in SETS:
        doc.add_heading(entry["name"], level=2)
        add_text(doc, f"Set ID: {SET_IDS[entry['name']]}", after=1, color=MUTED, size=9.5)
        meta = add_text(doc, f"融合方式：{entry['mode']}。完成后：{entry['unlock']}。", after=5, color=MUTED, size=10)
        for name, card_type, text in entry["cards"]:
            add_card(doc, name, card_type, text)
        result_name, result_type, result_lines = entry["result"]
        add_result(doc, result_name, result_type, result_lines)


def add_independents(doc):
    doc.add_heading("6. 独立军略", level=1)
    add_text(doc, "独立军略不属于任何 Set，不参与融合；每张本局最多持有 1 次。", after=4, color=MUTED, size=10)
    for name, card_type, text, eligibility in INDEPENDENTS:
        add_card(doc, name, card_type, text, eligibility)


def add_implementation_notes(doc):
    doc.add_heading("7. 配置与实现备注", level=1)
    add_label_value(doc, [
        ("推荐配置主键", "command_card_id、command_set_id、fusion_result_id、unlock_set_ids、offer_eligibility_rule、command_target_mode、target_priority_rule、effect_payload。"),
        ("运行时状态", "已购卡、已备战主动军令、已完成 Set、已解锁路线、军略图鉴状态、本波军令点、主动军令本波已用状态、当前小兵数量与补员次数。"),
        ("UI 事件", "选路预览、自动锁定目标、预警开始/结束、军令返还、Set 完成、主动进阶、被动吞噬、新解锁路线、补员成功/失败。"),
        ("待数值核算", "初始军资、波次收入、基础卡价、刷新曲线、报价权重、属性 Modifier 的叠加与舍入、防御与暴击公式。"),
    ])
    doc.add_heading("8. 程序验收重点", level=1)
    for item in [
        "主动军令不会改变英雄自动技能的释放逻辑，也不会把英雄主动技能加入军令库。",
        "所有路目标按战前归属结算；单位跨路移动后仍应被原路军令命中。",
        "根系与二阶 Set 完成时自动融合，已购组件收益不丢失；高阶主动保留原 HUD 槽位。",
        "小兵死亡跨波次保持损失；补员只能在备战发生，且不超过升星带兵上限。",
        "灼烧不叠加；火媒、断索、风助等触发需遵守各自的每波、每路或每单位限制。",
        "目标预警期间失效时，只对明确声明返还的军令返还军令点并恢复本波使用资格。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("9. 卡牌 ID 索引", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 4300, 2560])
    for i, title in enumerate(["名称", "command_card_id", "所属 Set"]):
        set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])
    for entry in SETS:
        for name, _, _ in entry["cards"]:
            row = table.add_row().cells
            values = [name, CARD_IDS[name], SET_IDS[entry["name"]]]
            for i, value in enumerate(values):
                p = row[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(value)
                set_run_font(r, size=8.8)
        result_name, _, _ = entry["result"]
        row = table.add_row().cells
        values = [result_name, CARD_IDS[result_name], SET_IDS[entry["name"]]]
        for i, value in enumerate(values):
            p = row[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=8.8, color=DARK_BLUE)
    for name, _, _, _ in INDEPENDENTS:
        row = table.add_row().cells
        values = [name, CARD_IDS[name], "independent"]
        for i, value in enumerate(values):
            p = row[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=8.8)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_global_rules(doc)
    add_set_overview(doc)
    add_sets(doc)
    add_independents(doc)
    add_implementation_notes(doc)
    doc.core_properties.title = "军令效果设计清单 V1"
    doc.core_properties.subject = "三国斗阵军令卡程序实现参考"
    doc.core_properties.author = "project01 主策划"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
