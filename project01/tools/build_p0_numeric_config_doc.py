from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"D:\codex\project01\deliverables\三国斗阵_P0数值配置手册_v0.1.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6472"
HEADER_FILL = "E8EEF5"
NOTE_FILL = "F4F6F9"
WHITE = "FFFFFF"
GRAY_BORDER = "C7D1DD"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph_format(p, before=0, after=6, line=1.25, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(doc, text, size=10.5, color=INK, bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, before, after, 1.25, align)
    r = p.add_run(text)
    set_font(r, size, color, bold, italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_format(p, 18, 10, 1.0)
        r = p.add_run(text)
        set_font(r, 16, BLUE, True)
    elif level == 2:
        set_paragraph_format(p, 14, 7, 1.0)
        r = p.add_run(text)
        set_font(r, 13, BLUE, True)
    else:
        set_paragraph_format(p, 10, 5, 1.0)
        r = p.add_run(text)
        set_font(r, 11.5, DARK_BLUE, True)
    p.style = f"Heading {min(level,3)}"
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_FILL)
    p = cell.paragraphs[0]
    set_paragraph_format(p, 0, 0, 1.2)
    r = p.add_run(f"{title}：")
    set_font(r, 10.5, DARK_BLUE, True)
    r = p.add_run(body)
    set_font(r, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    head = table.rows[0]
    set_repeat_table_header(head)
    for i, h in enumerate(headers):
        cell = head.cells[i]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        set_paragraph_format(p, 0, 0, 1.1, WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(str(h))
        set_font(r, font_size, DARK_BLUE, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(str(value)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p, 0, 0, 1.15, align)
            r = p.add_run(str(value))
            set_font(r, font_size, INK)
    set_table_geometry(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_check_table(doc, rows):
    return add_table(
        doc,
        ["编辑器项目", "填写 / 检查内容", "状态"],
        rows,
        [2100, 5500, 1760],
        9.3,
    )


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, 9, MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, DARK_BLUE)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(hp, 0, 0, 1.0)
    hr = hp.add_run("三国斗阵 | P0 数值配置手册")
    set_font(hr, 9, MUTED, True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(fp, 0, 0, 1.0)
    fr = fp.add_run("内部工作稿  |  第 ")
    set_font(fr, 9, MUTED)
    add_page_number(fp)
    fr = fp.add_run(" 页")
    set_font(fr, 9, MUTED)

    # Title block: compact_reference_guide + memo_masthead pattern.
    add_text(doc, "数值配置手册", 9.5, MUTED, True, before=0, after=8)
    title = doc.add_paragraph()
    set_paragraph_format(title, 0, 4, 1.0)
    tr = title.add_run("三国斗阵 P0 数值配置手册")
    set_font(tr, 24, "0B2545", True)
    add_text(doc, "用于 Y3 编辑器首轮属性、伤害、单位、难度与关卡 1 战斗样本配置", 12, MUTED, False, after=14)
    add_table(doc,
              ["版本", "状态", "适用范围"],
              [["v0.1", "已确认规则 + 首轮实测候选", "P0 首轮原型；不含装备、羁绊、完整军令经济与法术/真实伤害实测"]],
              [1500, 2700, 5160], 9.3)
    add_note(doc, "使用原则", "先按本手册配置项目级公式与基础单位；首关三波样本用于校准。标记为“候选 / 临时”的内容须在实测后再锁定，不应当作最终版本。")

    add_heading(doc, "1. 配置总览", 1)
    add_table(doc,
              ["模块", "当前规则", "编辑器落点"],
              [
                  ["物理伤害", "采用软减伤；护甲常数 K = 100", "项目设置 > 伤害公式"],
                  ["护甲数值", "整数；不采用默认数千到上万的面板尺度", "单位属性 / 属性成长"],
                  ["局内等级", "1 至 6 级；击杀 NPC 获得全图共享经验；升级仅增长属性", "英雄经验、单位击杀经验、属性成长"],
                  ["英雄技能", "固定 1 级；星级只解锁技能", "英雄技能配置、英雄星级"],
                  ["普攻距离", "近战 300；远程 600", "单位攻击范围"],
                  ["难度成长", "只可调血量、攻击、护甲、攻速、自然回血", "关卡触发 / 难度系数表"],
              ], [2100, 4050, 3210], 9.3)

    add_heading(doc, "2. 项目级战斗与属性设置", 1)
    add_heading(doc, "2.1 物理伤害公式", 2)
    add_text(doc, "项目采用以下物理伤害结算。有效护甲先结算穿透、减防等效果，最低为 0。", 10.5)
    add_note(doc, "Y3 伤害公式填写值", "伤害来源 伤害 * (1 - (伤害目标 防御(计算过穿透)) / (伤害目标 防御(计算过穿透) + 100)) * (1 + 伤害来源 所有伤害加成(%) * 0.01 - 伤害目标 伤害减免(%) * 0.01)")
    add_table(doc,
              ["有效护甲", "承受物理伤害", "物理减伤", "配置定位"],
              [["0", "100.0%", "0.0%", "无护甲"], ["25", "80.0%", "20.0%", "成型战士"], ["50", "66.7%", "33.3%", "6 级主坦锚点"], ["100", "50.0%", "50.0%", "高护甲检查点"], ["200", "33.3%", "66.7%", "P0 常规单位不应达到"]],
              [1500, 2500, 2200, 3160], 9.2)
    add_note(doc, "尚未确认", "法术伤害是否读取法术防御或复用本公式，以及真实伤害是否受“所有伤害加成 / 伤害减免”影响，均需另行实测。本手册不得以物理规则外推。")

    add_heading(doc, "2.2 通用属性规则", 2)
    add_table(doc,
              ["属性 / 系统", "确定规则"],
              [
                  ["攻击距离", "近战普攻模板 = 300；远程普攻模板 = 600。单位须先选定模板，偏离基准须另写原因。"],
                  ["单位升级", "最终属性 = 1 级基础属性 + (等级 - 1) × 属性成长。P0 局内上限为 6 级。"],
                  ["复合属性", "力量、敏捷、智力可映射至一个或多个一级属性。P0 不锁定映射系数，且正常升级不自动获得复合属性。"],
                  ["属性恢复", "项目级恢复频率为 1 秒；“生命恢复”按每秒 X 点理解。"],
                  ["技能等级", "skill_rank 固定为 1；局内等级与星级都不提高技能数值。"],
                  ["技能解锁", "英雄 1 / 2 / 3 / 4 / 5 / 6 星依次解锁主动、被动、主动、被动、主动、被动。"],
              ], [2300, 7060], 9.3)

    add_heading(doc, "3. 英雄基础面板（P0 候选）", 1)
    add_text(doc, "以下为 10 名英雄的 1 级基础值与每级成长。护甲已按 K=100 重标为整数；6 级数值按“基础 + 5 × 成长”计算。", 10.5)
    add_table(doc,
              ["英雄", "定位", "HP", "HP成长", "攻击", "攻成长", "护甲", "甲成长", "攻击间隔", "6级 HP / 攻击 / 护甲"],
              [
                  ["吕布", "前排压制", "3300", "340", "250", "25", "20", "2", "1.35", "5000 / 375 / 30"],
                  ["张飞", "主坦", "4000", "460", "170", "15", "30", "4", "1.45", "6300 / 245 / 50"],
                  ["吕蒙", "侧翼刺客", "2600", "280", "235", "26", "12", "2", "1.15", "4000 / 365 / 22"],
                  ["黄忠", "后排单点", "2300", "230", "255", "28", "9", "1", "1.40", "3450 / 395 / 14"],
                  ["诸葛亮", "范围控制", "2200", "250", "195", "22", "10", "1", "1.45", "3450 / 305 / 15"],
                  ["赵云", "机动战士", "3000", "320", "220", "24", "16", "2", "1.25", "4600 / 340 / 26"],
                  ["刘备", "治疗护盾", "2800", "350", "155", "15", "15", "2", "1.50", "4550 / 230 / 25"],
                  ["夏侯惇", "主坦反斩", "3800", "440", "180", "16", "32", "4", "1.45", "6000 / 260 / 52"],
                  ["司马懿", "后排压制", "2100", "230", "205", "25", "9", "1", "1.45", "3250 / 330 / 14"],
                  ["甘宁", "侧翼爆发", "2750", "290", "240", "27", "14", "2", "1.15", "4200 / 375 / 24"],
              ], [700, 1050, 650, 690, 650, 690, 590, 650, 820, 2870], 7.6)
    add_note(doc, "英雄射程", "每名英雄的近战 / 远程普攻类型仍应按角色战斗类型落入单位物编；填写后分别使用 300 / 600 基准。本表不预先虚构未定的英雄射程。")

    add_heading(doc, "4. NPC 基础面板与敌方英雄规则", 1)
    add_heading(doc, "4.1 原生 NPC（难度 1）", 2)
    add_table(doc,
              ["单位", "身份", "最大生命", "物理攻击", "物理护甲", "攻击间隔", "普攻距离", "P0 行为"],
              [
                  ["步卒", "近战小兵", "850", "72", "6", "1.65", "300", "仅近战普攻；击杀经验 10"],
                  ["弓手", "远程小兵", "650", "88", "3", "1.75", "600", "仅远程普攻；击杀经验 18"],
                  ["力士", "近战队长", "3200", "150", "26", "1.85", "300", "普攻为主；击杀经验 35；重击待后续启用"],
              ], [900, 1000, 1050, 1000, 900, 1050, 950, 2510], 8.8)
    add_text(doc, "P0 原生 NPC 仅为步卒、弓手、力士。力士只能作为队长，不能作为小兵；力士死亡后，所属 NPC 小兵继续战斗。", 10.2)

    add_heading(doc, "4.2 敌方三国英雄队长", 2)
    add_table(doc,
              ["字段", "填写规则"],
              [
                  ["hero_id", "复用已设计三国英雄的基础面板、成长、技能与表现资产。"],
                  ["npc_variant_id", "只覆写敌方初始站位、小兵模板和数量；不得提供额外属性倍率或削减技能数量。"],
                  ["npc_level", "固定 1 至 6 级；决定基础属性成长；战斗中不获得经验。"],
                  ["npc_star_level", "固定 1 至 6 星；自动携带所有解锁星级不高于该值的技能。6 星即完整 6 技能。"],
                  ["soldier_template_id / soldier_count", "配置该敌方队长的小兵模板与数量。"],
              ], [2600, 6760], 9.3)
    add_note(doc, "敌方满技能限制", "难度 9 至 10 可出现 6 星满技能敌方英雄；每波最多 1 名，波次预览必须显示星级与关键技能。")

    add_heading(doc, "5. 关卡难度属性系数", 1)
    add_text(doc, "难度可调属性池只允许：最大生命、物理攻击、物理护甲、攻击速度、自然回血。P0 首轮只启用前三项；攻击速度与自然回血字段保留但固定为 1.0。", 10.5)
    add_note(doc, "生成顺序", "队长等级基础属性 = 基础属性 + (npc_level - 1) × 属性成长；最终 NPC 属性 = round(队长等级基础属性 × 关卡难度属性系数)。步卒、弓手没有等级成长，直接套难度系数。")
    add_table(doc,
              ["难度", "血量系数", "攻击系数", "护甲系数", "攻速系数", "自然回血系数"],
              [
                  ["1", "1.00", "1.00", "1.00", "1.00", "1.00"],
                  ["2", "1.12", "1.06", "1.01", "1.00", "1.00"],
                  ["3", "1.26", "1.12", "1.03", "1.00", "1.00"],
                  ["4", "1.42", "1.18", "1.05", "1.00", "1.00"],
                  ["5", "1.60", "1.25", "1.08", "1.00", "1.00"],
                  ["6", "1.82", "1.33", "1.10", "1.00", "1.00"],
                  ["7", "2.08", "1.42", "1.13", "1.00", "1.00"],
                  ["8", "2.38", "1.52", "1.17", "1.00", "1.00"],
                  ["9", "2.72", "1.63", "1.21", "1.00", "1.00"],
                  ["10", "3.10", "1.75", "1.26", "1.00", "1.00"],
              ], [900, 1700, 1700, 1700, 1680, 1680], 9.1)
    add_text(doc, "攻击范围、命中、闪避、暴击、控制时长和技能 CD 不参与关卡难度自动缩放。", 10.2)

    add_heading(doc, "6. 英雄经验与技能数值基线", 1)
    add_heading(doc, "6.1 局内经验表（候选）", 2)
    add_table(doc,
              ["升级", "所需经验", "累计经验"],
              [["1 → 2", "100", "100"], ["2 → 3", "150", "250"], ["3 → 4", "225", "475"], ["4 → 5", "325", "800"], ["5 → 6", "450", "1250"]],
              [2500, 3100, 3760], 9.5)
    add_text(doc, "启用全图经验：每名敌方单位死亡时，按其击杀经验同时给予所有上阵英雄；单局结束后等级与经验清空。波次默认不发经验，只有特殊事件、关卡奖励或技能机制才可配置额外的 wave_clear_xp_bonus（默认 0）。经验曲线为候选，待完整 10 至 12 波关卡反推后确认。", 10.2)
    add_heading(doc, "6.2 主动技能统一档位（供后续逐英雄落数值）", 2)
    add_table(doc,
              ["项目", "P0 推荐值"],
              [
                  ["阶段 1 主动 CD", "8 秒（允许 7 至 9 秒）"],
                  ["阶段 2 主动 CD", "11 秒（允许 10 至 12 秒）"],
                  ["阶段 3 主动 CD", "15 秒（允许 13 至 16 秒）"],
                  ["阶段 1 单体伤害", "160% 至 220% 攻击"],
                  ["阶段 2 范围伤害", "每目标 100% 至 150% 攻击"],
                  ["阶段 3 单体爆发", "280% 至 360% 攻击"],
                  ["单体治疗 / 护盾", "目标最大生命 12% 至 18% / 14% 至 22%"],
              ], [3100, 6260], 9.3)

    add_heading(doc, "7. 关卡 1 · 难度 1 三波战斗样本（实测候选）", 1)
    add_note(doc, "用途", "这是基础战斗校准样本，不是正式新手开局、完整波次表或剧情关。它不测试军令市场、补员、Boss、敌方英雄或高难度系数。")
    add_heading(doc, "7.1 我方固定配置", 2)
    add_table(doc,
              ["项目", "配置"],
              [
                  ["队长 A", "张飞，1 星，1 级，格位 2（前排中路）；带 2 名近战小兵。"],
                  ["队长 B", "黄忠，1 星，1 级，格位 8（后排中路）；带 2 名远程小兵。"],
                  ["局外 / 军令", "无局外属性加成、无军令、无补员。"],
                  ["力士重击", "关闭；首轮只校准普攻、队长技能与站位。"],
              ], [2300, 7060], 9.3)
    add_table(doc,
              ["P0 统一友方小兵", "最大生命", "物理攻击", "物理护甲", "攻击间隔", "普攻距离"],
              [["近战小兵", "600", "55", "5", "1.70", "300"], ["远程小兵", "450", "65", "3", "1.80", "600"]],
              [2200, 1500, 1400, 1400, 1500, 1360], 9.0)
    add_heading(doc, "7.2 本样本的技能临时覆写", 2)
    add_table(doc,
              ["英雄", "技能", "临时数值", "自动规则"],
              [
                  ["张飞", "当阳怒吼（1 星）", "180% 物理攻击伤害；嘲讽 3 秒；CD 8 秒。", "优先敌方队长；无队长时当前目标。"],
                  ["黄忠", "定军箭（1 星）", "190% 物理攻击伤害；攻速 -20%，持续 5 秒；CD 8 秒。", "当前目标。"],
              ], [900, 1700, 3600, 3160], 8.8)
    add_text(doc, "上述两项只作为本样本 TestSkillOverride，不能替代英雄技能总表的最终审核。", 10.0, MUTED, italic=True)
    add_heading(doc, "7.3 波次、经验与验收", 2)
    add_table(doc,
              ["波次", "敌方配置（均在中路 2）", "击杀经验构成", "目标时长", "验证点"],
              [
                  ["W1", "力士 ×1 + 步卒 ×1", "35 + 10 = 45", "15–18 秒", "基础接敌、近战承伤、后排输出。"],
                  ["W2", "力士 ×1 + 步卒 ×2", "35 + 10×2 = 55", "17–22 秒", "张飞承伤与嘲讽；波后升至 2 级。"],
                  ["W3", "力士 ×1 + 步卒 ×2 + 弓手 ×2", "35 + 10×2 + 18×2 = 91", "20–25 秒", "前排保护、远程压力、队长死后残兵清理。"],
              ], [800, 2700, 900, 1300, 3660], 8.8)
    add_table(doc,
              ["验收项", "通过标准"],
              [
                  ["通关", "三波全部胜利；张飞、黄忠均存活。"],
                  ["承伤", "W3 结束时张飞至少保留 15% 生命；黄忠不应承受主要近战伤害。"],
                  ["小兵", "允许 W3 阵亡 0 至 2 名友方小兵；全灭时须记录原因。"],
                  ["记录", "每波耗时、双方队长最终生命、小兵伤亡、技能释放次数与第一次释放时点。"],
              ], [1900, 7460], 9.2)

    add_heading(doc, "8. 编辑器配置清单", 1)
    add_check_table(doc,
                    [
                        ["物理伤害公式", "将默认公式分母中的 +10000 改为 +100，并确认项目实际生效。", "必须"],
                        ["项目属性恢复频率", "设置为 1 秒。", "必须"],
                        ["普攻攻击距离", "近战模板 300；远程模板 600。", "必须"],
                        ["英雄基础属性", "录入第 3 节 10 名英雄的 HP、攻击、护甲、成长和攻击间隔。", "必须"],
                        ["NPC 基础属性", "录入步卒、弓手、力士；力士设为队长，步卒/弓手不设为队长。", "必须"],
                        ["难度系数", "录入第 5 节；P0 攻速 / 自然回血均为 1.0。", "必须"],
                        ["经验列表与单位经验", "录入第 6.1 节 1 至 6 级经验表；开启全图经验；步卒/弓手/力士击杀经验分别为 10/18/35。", "候选"],
                        ["英雄带兵上限", "1 至 6 星小兵上限依次为 2/3/4/5/6/8；实际数量不得超限。", "必须"],
                        ["英雄小兵类型", "按第 9 节为每名英雄配置单一近战或远程小兵模板。", "必须"],
                        ["我方小兵面板", "近战 600/55/5/1.70/300；远程 450/65/3/1.80/600；小兵无局内等级成长。", "候选"],
                        ["三波样本", "按第 7 节配置张飞、黄忠、友方小兵、敌方波次与临时技能覆写。", "实测"],
                        ["法术 / 真实伤害", "暂不按本手册外推；完成独立实测后再配置。", "暂缓"],
                    ])
    add_note(doc, "调参优先级", "实测异常时先检查物理公式、攻击距离、攻击间隔、站位与难度系数；随后调整敌方小兵数量。不要在首轮就用攻速、自然回血、命中、闪避、暴击、控制时长或 CD 缩放掩盖基础问题。")

    add_heading(doc, "9. 英雄小队与带兵规则", 1)
    add_text(doc, "英雄升星同时解锁技能并提高可携带小兵上限；P0 不提供小兵等级成长、混编兵种、第三兵种或独立小兵技能树。", 10.5)
    add_table(doc,
              ["英雄星级", "小兵上限", "说明"],
              [["1 星", "2", "基础小队；与关卡 1 三波样本一致。"], ["2 星", "3", "第一项被动后的基础兵力。"], ["3 星", "4", "第二主动与更多小兵并行。"], ["4 星", "5", "小队被动开始稳定放大兵力价值。"], ["5 星", "6", "高星主动与更完整队列。"], ["6 星", "8", "满技能的队伍规模奖励；唯一一次 +2 跃升。"]],
              [1800, 1800, 5760], 9.3)
    add_table(doc,
              ["英雄", "默认小兵", "英雄", "默认小兵"],
              [["吕布", "近战", "张飞", "近战"], ["吕蒙", "近战", "黄忠", "远程"], ["诸葛亮", "远程", "赵云", "近战"], ["刘备", "远程", "夏侯惇", "近战"], ["司马懿", "远程", "甘宁", "近战"]],
              [2000, 2680, 2000, 2680], 9.2)
    add_table(doc,
              ["我方小兵", "最大生命", "物理攻击", "物理护甲", "攻击间隔", "普攻距离", "局内成长"],
              [["近战小兵", "600", "55", "5", "1.70", "300", "无"], ["远程小兵", "450", "65", "3", "1.80", "600", "无"]],
              [1800, 1300, 1300, 1300, 1400, 1400, 860], 8.8)
    add_text(doc, "敌方英雄队长使用同一星级带兵上限；npc_star_level 同时决定技能解锁和上限。波次可填写低于上限的实际数量，但不得超限；首轮不允许一名队长携带近战与远程混编。", 10.2)

    doc.core_properties.title = "三国斗阵 P0 数值配置手册"
    doc.core_properties.subject = "Y3 编辑器首轮数值配置"
    doc.core_properties.author = "项目策划"
    doc.core_properties.comments = "P0 v0.1"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
